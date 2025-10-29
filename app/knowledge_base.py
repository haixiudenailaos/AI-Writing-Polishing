"""
知识库管理模块
负责文件扫描、向量化处理和存储
"""

import os
import json
import hashlib
from typing import List, Dict, Any, Optional, Callable
from pathlib import Path
from dataclasses import dataclass, asdict
from datetime import datetime
import requests


@dataclass
class VectorDocument:
    """向量化文档"""
    id: str
    file_path: str
    content: str
    vector: List[float]
    metadata: Dict[str, Any]
    created_at: str


@dataclass
class KnowledgeBase:
    """知识库"""
    id: str
    name: str
    root_path: str
    documents: List[VectorDocument]
    created_at: str
    updated_at: str
    metadata: Dict[str, Any]


class VectorEmbeddingClient:
    """阿里云向量化客户端
    
    支持的模型：
    - text-embedding-v4: 阿里云通义千问嵌入模型，支持多语言，输出维度1024
    
    支持两种API调用方式：
    1. 兼容OpenAI格式（默认）：https://dashscope.aliyuncs.com/compatible-mode/v1/embeddings
    2. 原生DashScope API：https://dashscope.aliyuncs.com/api/v1/services/embeddings/text-embedding/text-embedding
    
    参考文档：
    - https://help.aliyun.com/zh/model-studio/developer-reference/text-embedding-v3-api
    """
    
    def __init__(self, api_key: str, model: str = "text-embedding-v4", use_native_api: bool = False):
        """
        初始化向量化客户端
        
        Args:
            api_key: 阿里云API密钥（DashScope API Key）
            model: 向量模型名称，默认 text-embedding-v4
            use_native_api: 是否使用原生DashScope API（默认使用兼容OpenAI格式）
        """
        self.api_key = api_key
        self.model = model
        self.use_native_api = use_native_api
        
        # 根据选择设置endpoint
        if use_native_api:
            self.base_url = "https://dashscope.aliyuncs.com/api/v1/services/embeddings/text-embedding/text-embedding"
        else:
            # 兼容OpenAI格式的endpoint（推荐，更通用）
            self.base_url = "https://dashscope.aliyuncs.com/compatible-mode/v1/embeddings"
    
    def embed_text(self, text: str) -> List[float]:
        """
        对文本进行向量化
        
        Args:
            text: 待向量化的文本
            
        Returns:
            向量列表（1024维）
            
        Raises:
            ValueError: 文本为空
            RuntimeError: API调用失败
        """
        if not text or not text.strip():
            raise ValueError("文本内容不能为空")
        
        # 限制文本长度，避免超出模型限制（text-embedding-v4支持最长2048 tokens）
        if len(text) > 6000:  # 粗略估计，约2000个中文字符
            text = text[:6000]
        
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        
        # 根据API类型构造不同的payload
        if self.use_native_api:
            # 原生DashScope API格式
            payload = {
                "model": self.model,
                "input": {
                    "texts": [text]
                }
            }
        else:
            # 兼容OpenAI格式
            payload = {
                "model": self.model,
                "input": text,
                "encoding_format": "float"
            }
        
        try:
            response = requests.post(
                self.base_url,
                headers=headers,
                json=payload,
                timeout=30
            )
            
            # 详细的错误处理
            if response.status_code == 401:
                raise RuntimeError("API密钥认证失败，请检查阿里云DashScope API密钥是否正确")
            
            if response.status_code == 400:
                error_detail = response.json().get("message", response.text)
                raise RuntimeError(f"请求参数错误: {error_detail}")
            
            if response.status_code == 429:
                raise RuntimeError("API调用频率超限，请稍后重试")
            
            if not response.ok:
                raise RuntimeError(f"向量化请求失败: HTTP {response.status_code} - {response.text}")
            
            data = response.json()
            
            # 根据API类型解析不同的响应格式
            if self.use_native_api:
                # 原生API格式: {"output": {"embeddings": [{"embedding": [...], "text_index": 0}]}}
                if "output" not in data or "embeddings" not in data["output"]:
                    raise RuntimeError(f"向量化响应格式错误: {data}")
                
                embeddings = data["output"]["embeddings"]
                if not embeddings or len(embeddings) == 0:
                    raise RuntimeError("向量化结果为空")
                
                embedding = embeddings[0].get("embedding")
            else:
                # 兼容OpenAI格式: {"data": [{"embedding": [...]}]}
                if "data" not in data or not data["data"]:
                    raise RuntimeError(f"向量化响应格式错误: {data}")
                
                embedding = data["data"][0].get("embedding")
            
            if not embedding:
                raise RuntimeError("向量化结果为空")
            
            # 验证向量维度（text-embedding-v4应该是1024维）
            if len(embedding) != 1024:
                print(f"[WARN] 向量维度异常: 预期1024维，实际{len(embedding)}维")
            
            return embedding
            
        except requests.Timeout:
            raise RuntimeError("网络请求超时，请检查网络连接")
        except requests.RequestException as e:
            raise RuntimeError(f"网络请求失败: {str(e)}")
    
    def embed_batch(self, texts: List[str], 
                    progress_callback: Optional[Callable[[int, int, str], None]] = None,
                    batch_size: int = 10) -> List[List[float]]:
        """
        批量向量化文本（支持批量API调用以提高效率）
        
        Args:
            texts: 文本列表
            progress_callback: 进度回调函数 (current, total, message)
            batch_size: 每批处理的文本数量，阿里云API限制最大10个
            
        Returns:
            向量列表
        """
        vectors = []
        total = len(texts)
        
        # 按批次处理
        for batch_start in range(0, total, batch_size):
            batch_end = min(batch_start + batch_size, total)
            batch_texts = texts[batch_start:batch_end]
            
            try:
                if progress_callback:
                    progress_callback(
                        batch_end, total, 
                        f"正在处理第 {batch_start + 1}-{batch_end}/{total} 个文本..."
                    )
                
                # 批量调用API
                batch_vectors = self._embed_batch_api(batch_texts)
                vectors.extend(batch_vectors)
                
            except Exception as e:
                if progress_callback:
                    progress_callback(batch_end, total, f"批次处理失败: {str(e)}，尝试逐个处理...")
                
                # 失败时逐个处理该批次
                for text in batch_texts:
                    try:
                        vector = self.embed_text(text)
                        vectors.append(vector)
                    except Exception as e2:
                        print(f"[ERROR] 单个文本向量化失败: {str(e2)}")
                        vectors.append([])  # 失败时添加空向量
        
        return vectors
    
    def _embed_batch_api(self, texts: List[str]) -> List[List[float]]:
        """
        批量API调用（text-embedding-v4支持批量输入）
        
        Args:
            texts: 文本列表
            
        Returns:
            向量列表
        """
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        
        # 根据API类型构造不同的payload
        if self.use_native_api:
            # 原生DashScope API格式
            payload = {
                "model": self.model,
                "input": {
                    "texts": texts
                }
            }
        else:
            # 兼容OpenAI格式（支持列表输入）
            payload = {
                "model": self.model,
                "input": texts,
                "encoding_format": "float"
            }
        
        try:
            response = requests.post(
                self.base_url,
                headers=headers,
                json=payload,
                timeout=60  # 批量请求需要更长超时
            )
            
            if not response.ok:
                raise RuntimeError(f"批量向量化失败: {response.status_code} {response.text}")
            
            data = response.json()
            
            # 根据API类型解析不同的响应格式
            vectors = []
            if self.use_native_api:
                # 原生API格式
                if "output" not in data or "embeddings" not in data["output"]:
                    raise RuntimeError("响应格式错误")
                
                embeddings = data["output"]["embeddings"]
                # 按text_index排序
                embeddings.sort(key=lambda x: x.get("text_index", 0))
                for item in embeddings:
                    embedding = item.get("embedding")
                    vectors.append(embedding if embedding else [])
            else:
                # 兼容OpenAI格式
                if "data" not in data:
                    raise RuntimeError("响应格式错误")
                
                for item in data["data"]:
                    embedding = item.get("embedding")
                    vectors.append(embedding if embedding else [])
            
            return vectors
            
        except Exception as e:
            raise RuntimeError(f"批量API调用失败: {str(e)}")


class RerankClient:
    """阿里云重排序客户端
    
    用于对检索结果进行重新排序，提升相关性
    
    支持的模型：
    - gte-rerank-v2: 通用文本重排序模型
    - qwen3-rerank: 通义千问3重排序模型
    
    参考文档：
    - https://help.aliyun.com/zh/model-studio/developer-reference/rerank-api
    """
    
    def __init__(self, api_key: str, model: str = "gte-rerank-v2"):
        """
        初始化重排序客户端
        
        Args:
            api_key: 阿里云API密钥（DashScope API Key）
            model: 重排序模型名称
        """
        self.api_key = api_key
        self.model = model
        self.base_url = "https://dashscope.aliyuncs.com/api/v1/services/rerank/text-rerank/text-rerank"
    
    def rerank(self, query: str, documents: List[str], top_n: int = 5) -> List[Dict[str, Any]]:
        """
        对文档列表进行重排序
        
        Args:
            query: 查询文本
            documents: 候选文档列表
            top_n: 返回前N个最相关的文档
            
        Returns:
            重排序后的文档列表，每个元素包含 {index, document, relevance_score}
        """
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        
        payload = {
            "model": self.model,
            "input": {
                "query": query,
                "documents": documents
            },
            "parameters": {
                "return_documents": True,
                "top_n": min(top_n, len(documents))
            }
        }
        
        try:
            response = requests.post(
                self.base_url,
                headers=headers,
                json=payload,
                timeout=30
            )
            
            if response.status_code == 401:
                raise RuntimeError("API密钥认证失败")
            
            if not response.ok:
                raise RuntimeError(f"重排序请求失败: {response.status_code} {response.text}")
            
            data = response.json()
            
            # 解析响应
            if "output" not in data or "results" not in data["output"]:
                raise RuntimeError("重排序响应格式错误")
            
            results = data["output"]["results"]
            
            # 构造返回结果
            reranked = []
            for item in results:
                reranked.append({
                    "index": item.get("index"),
                    "document": item.get("document"),
                    "relevance_score": item.get("relevance_score", 0.0)
                })
            
            return reranked
            
        except requests.RequestException as e:
            raise RuntimeError(f"重排序网络请求失败: {str(e)}")


class KnowledgeBaseManager:
    """知识库管理器"""
    
    SUPPORTED_EXTENSIONS = {'.txt', '.md', '.docx', '.epub', '.doc', '.rtf', '.odt'}
    
    def __init__(self, storage_dir: str = "app_data/knowledge_bases"):
        """
        初始化知识库管理器
        
        Args:
            storage_dir: 知识库存储目录
        """
        # 确保使用绝对路径，避免路径问题
        self.storage_dir = Path(storage_dir).resolve()
        self.storage_dir.mkdir(parents=True, exist_ok=True)
        
        self.embedding_client: Optional[VectorEmbeddingClient] = None
    
    def set_embedding_client(self, api_key: str, model: str = "text-embedding-v4"):
        """设置向量化客户端"""
        self.embedding_client = VectorEmbeddingClient(api_key, model)
    
    def test_embedding_connection(self) -> tuple[bool, str]:
        """
        测试向量化API连接
        
        Returns:
            (是否成功, 消息)
        """
        if not self.embedding_client:
            return False, "未配置向量化客户端"
        
        try:
            # 使用简单的测试文本
            test_text = "测试向量化API连接"
            vector = self.embedding_client.embed_text(test_text)
            
            if vector and len(vector) == 1024:
                return True, f"连接成功！向量维度: {len(vector)}"
            else:
                return False, f"向量维度异常: {len(vector) if vector else 0}"
                
        except Exception as e:
            return False, f"连接失败: {str(e)}"
    
    def scan_folder(self, folder_path: str, 
                    progress_callback: Optional[Callable[[int, int, str], None]] = None) -> List[str]:
        """
        扫描文件夹，查找所有支持的文件
        
        Args:
            folder_path: 文件夹路径
            progress_callback: 进度回调函数
            
        Returns:
            文件路径列表
        """
        folder = Path(folder_path)
        if not folder.exists() or not folder.is_dir():
            raise ValueError(f"无效的文件夹路径: {folder_path}")
        
        files = []
        
        # 递归扫描
        for root, _, filenames in os.walk(folder):
            for filename in filenames:
                file_path = Path(root) / filename
                if file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                    files.append(str(file_path))
                    
                    if progress_callback:
                        progress_callback(
                            len(files), 0, 
                            f"已发现 {len(files)} 个文件..."
                        )
        
        return files
    
    def read_file_content(self, file_path: str) -> str:
        """
        读取文件内容
        
        Args:
            file_path: 文件路径
            
        Returns:
            文件内容
        """
        path = Path(file_path)
        suffix = path.suffix.lower()
        
        try:
            if suffix == '.txt' or suffix == '.md':
                with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
            
            elif suffix == '.docx':
                # 需要python-docx库
                try:
                    from docx import Document
                    doc = Document(path)
                    return '\n'.join([p.text for p in doc.paragraphs])
                except ImportError:
                    return f"[无法读取 .docx 文件，需要安装 python-docx 库]"
            
            elif suffix == '.epub':
                # 需要ebooklib库
                try:
                    import ebooklib
                    from ebooklib import epub
                    from bs4 import BeautifulSoup
                    
                    book = epub.read_epub(path)
                    content = []
                    for item in book.get_items():
                        if item.get_type() == ebooklib.ITEM_DOCUMENT:
                            soup = BeautifulSoup(item.get_content(), 'html.parser')
                            content.append(soup.get_text())
                    return '\n'.join(content)
                except ImportError:
                    return f"[无法读取 .epub 文件，需要安装 ebooklib 和 beautifulsoup4 库]"
            
            else:
                # 其他格式尝试文本读取
                with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
                    
        except Exception as e:
            return f"[读取文件失败: {str(e)}]"
    
    def create_knowledge_base(
        self, 
        name: str, 
        folder_path: str,
        chunk_size: int = 800,  # 增大分块大小以更好利用模型能力
        progress_callback: Optional[Callable[[int, int, str], None]] = None,
        error_callback: Optional[Callable[[str], None]] = None
    ) -> Optional[KnowledgeBase]:
        """
        创建知识库
        
        Args:
            name: 知识库名称
            folder_path: 源文件夹路径
            chunk_size: 文本分块大小（字符数）
            progress_callback: 进度回调
            error_callback: 错误回调
            
        Returns:
            知识库对象，失败时返回None
        """
        if not self.embedding_client:
            if error_callback:
                error_callback("未配置向量化客户端，请先配置阿里云API密钥")
            return None
        
        try:
            # 1. 扫描文件
            if progress_callback:
                progress_callback(0, 100, "正在扫描文件夹...")
            
            files = self.scan_folder(folder_path, progress_callback)
            
            if not files:
                if error_callback:
                    error_callback("文件夹中没有找到支持的文件")
                return None
            
            # 2. 读取和分块
            if progress_callback:
                progress_callback(10, 100, f"正在读取 {len(files)} 个文件...")
            
            chunks = []
            chunk_metadata = []
            
            for i, file_path in enumerate(files):
                content = self.read_file_content(file_path)
                
                # 分块
                file_chunks = self._chunk_text(content, chunk_size)
                
                for j, chunk in enumerate(file_chunks):
                    chunks.append(chunk)
                    chunk_metadata.append({
                        'file_path': file_path,
                        'chunk_index': j,
                        'total_chunks': len(file_chunks)
                    })
                
                if progress_callback:
                    progress = 10 + int((i + 1) / len(files) * 30)
                    progress_callback(
                        progress, 100, 
                        f"已读取 {i + 1}/{len(files)} 个文件，生成 {len(chunks)} 个文本块"
                    )
            
            # 3. 向量化（使用批量处理提高效率）
            if progress_callback:
                progress_callback(40, 100, "正在进行向量化处理...")
            
            def vector_progress(current, total, message):
                """向量化进度回调"""
                progress = 40 + int(current / total * 50)
                if progress_callback:
                    progress_callback(progress, 100, message)
            
            try:
                # 使用批量API提高效率
                vectors = self.embedding_client.embed_batch(
                    chunks, 
                    progress_callback=vector_progress,
                    batch_size=10  # 每批10个文本块（阿里云API限制）
                )
                
                # 检查向量化结果
                failed_count = sum(1 for v in vectors if not v)
                if failed_count > 0:
                    if error_callback:
                        error_callback(f"有 {failed_count} 个文本块向量化失败，已跳过")
                
            except Exception as e:
                if error_callback:
                    error_callback(f"批量向量化失败: {str(e)}")
                return None
            
            # 4. 创建文档对象（过滤掉向量化失败的文档）
            documents = []
            skipped_count = 0
            for i, (chunk, vector, metadata) in enumerate(zip(chunks, vectors, chunk_metadata)):
                # 跳过向量化失败的文档
                if not vector:
                    skipped_count += 1
                    continue
                
                doc_id = self._generate_doc_id(chunk, metadata['file_path'])
                doc = VectorDocument(
                    id=doc_id,
                    file_path=metadata['file_path'],
                    content=chunk,
                    vector=vector,
                    metadata=metadata,
                    created_at=datetime.now().isoformat()
                )
                documents.append(doc)
            
            if skipped_count > 0:
                print(f"[WARN] 跳过了 {skipped_count} 个向量化失败的文档")
            
            if not documents:
                if error_callback:
                    error_callback("所有文档向量化失败，无法创建知识库")
                return None
            
            # 5. 创建知识库
            kb_id = self._generate_kb_id(name)
            kb = KnowledgeBase(
                id=kb_id,
                name=name,
                root_path=folder_path,
                documents=documents,
                created_at=datetime.now().isoformat(),
                updated_at=datetime.now().isoformat(),
                metadata={
                    'total_files': len(files),
                    'total_chunks': len(chunks),
                    'chunk_size': chunk_size
                }
            )
            
            # 6. 保存知识库到用户选中的文件夹
            self._save_knowledge_base(kb, folder_path)
            
            if progress_callback:
                progress_callback(100, 100, "知识库创建完成！")
            
            return kb
            
        except Exception as e:
            if error_callback:
                error_callback(f"创建知识库失败: {str(e)}")
            return None
    
    def _chunk_text(self, text: str, chunk_size: int) -> List[str]:
        """将文本分块"""
        if not text:
            return []
        
        chunks = []
        lines = text.split('\n')
        current_chunk = []
        current_size = 0
        
        for line in lines:
            line_size = len(line)
            
            if current_size + line_size > chunk_size and current_chunk:
                chunks.append('\n'.join(current_chunk))
                current_chunk = [line]
                current_size = line_size
            else:
                current_chunk.append(line)
                current_size += line_size
        
        if current_chunk:
            chunks.append('\n'.join(current_chunk))
        
        return chunks
    
    def _generate_doc_id(self, content: str, file_path: str) -> str:
        """生成文档ID"""
        hash_input = f"{content}{file_path}".encode('utf-8')
        return hashlib.md5(hash_input).hexdigest()
    
    def _generate_kb_id(self, name: str) -> str:
        """生成知识库ID"""
        timestamp = datetime.now().isoformat()
        hash_input = f"{name}{timestamp}".encode('utf-8')
        return hashlib.md5(hash_input).hexdigest()
    
    def _save_knowledge_base(self, kb: KnowledgeBase, target_folder: Optional[str] = None):
        """保存知识库到指定文件夹
        
        Args:
            kb: 知识库对象
            target_folder: 目标文件夹路径，如果为None则使用默认存储目录
        """
        try:
            if target_folder:
                # 保存到用户指定的文件夹
                target_path = Path(target_folder)
                kb_file = target_path / f".knowledge_base_{kb.name}.json"
            else:
                # 保存到默认存储目录（向后兼容）
                kb_file = self.storage_dir / f"{kb.id}.json"
            
            # 转换为字典
            kb_dict = asdict(kb)
            
            # 确保目标目录存在
            kb_file.parent.mkdir(parents=True, exist_ok=True)
            
            # 保存到文件
            with open(kb_file, 'w', encoding='utf-8') as f:
                json.dump(kb_dict, f, ensure_ascii=False, indent=2)
            
            print(f"[INFO] 知识库已保存: {kb_file}")
            
            # 同时在默认目录保存一份引用（用于全局管理）
            if target_folder:
                reference_file = self.storage_dir / f"{kb.id}.json"
                self.storage_dir.mkdir(parents=True, exist_ok=True)
                
                # 保存引用信息
                reference_data = {
                    "id": kb.id,
                    "name": kb.name,
                    "root_path": kb.root_path,
                    "kb_file_path": str(kb_file.resolve()),
                    "created_at": kb.created_at,
                    "updated_at": kb.updated_at,
                    "metadata": kb.metadata
                }
                
                with open(reference_file, 'w', encoding='utf-8') as f:
                    json.dump(reference_data, f, ensure_ascii=False, indent=2)
            
        except Exception as e:
            raise RuntimeError(f"保存知识库失败: {str(e)}")
    
    def load_knowledge_base(self, kb_id: str) -> Optional[KnowledgeBase]:
        """加载知识库
        
        Args:
            kb_id: 知识库ID
            
        Returns:
            知识库对象，如果不存在则返回None
        """
        reference_file = self.storage_dir / f"{kb_id}.json"
        
        if not reference_file.exists():
            return None
        
        with open(reference_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # 检查是否是引用文件
        if "kb_file_path" in data:
            # 这是一个引用，加载实际的知识库文件
            kb_file_path = Path(data["kb_file_path"])
            if not kb_file_path.exists():
                print(f"[WARN] 知识库文件不存在: {kb_file_path}")
                return None
            
            with open(kb_file_path, 'r', encoding='utf-8') as f:
                kb_dict = json.load(f)
        else:
            # 这是旧格式的知识库文件，直接加载
            kb_dict = data
        
        # 重构文档对象
        documents = [
            VectorDocument(**doc) for doc in kb_dict['documents']
        ]
        
        kb_dict['documents'] = documents
        return KnowledgeBase(**kb_dict)
    
    def list_knowledge_bases(self) -> List[Dict[str, Any]]:
        """列出所有知识库"""
        kb_list = []
        
        for kb_file in self.storage_dir.glob("*.json"):
            try:
                with open(kb_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                
                # 检查是否是引用文件
                if "kb_file_path" in data:
                    # 引用文件格式
                    kb_list.append({
                        'id': data['id'],
                        'name': data['name'],
                        'created_at': data['created_at'],
                        'root_path': data.get('root_path', ''),
                        'kb_file_path': data.get('kb_file_path', ''),
                        'total_documents': data.get('metadata', {}).get('total_chunks', 0)
                    })
                else:
                    # 旧格式知识库文件
                    kb_list.append({
                        'id': data['id'],
                        'name': data['name'],
                        'created_at': data['created_at'],
                        'root_path': data.get('root_path', ''),
                        'kb_file_path': str(kb_file),
                        'total_documents': len(data.get('documents', []))
                    })
            except Exception as e:
                print(f"[WARN] 读取知识库文件失败 {kb_file}: {e}")
                continue
        
        return kb_list
    
    def delete_knowledge_base(self, kb_id: str) -> bool:
        """删除知识库"""
        kb_file = self.storage_dir / f"{kb_id}.json"
        
        if kb_file.exists():
            kb_file.unlink()
            return True
        
        return False
