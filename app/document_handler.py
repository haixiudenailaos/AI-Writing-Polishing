"""
文档处理器
支持读写多种格式文档（txt, doc, docx, pdf, md, rtf, odt, html等）
"""

import os
import sys
import zipfile
import re
from typing import Optional, List, Tuple
from pathlib import Path
from xml.sax.saxutils import escape
import tempfile


class DocumentHandler:
    """文档处理器 - 支持多种格式"""
    
    @staticmethod
    def read_document(file_path: str) -> Optional[str]:
        """读取文档内容
        
        Args:
            file_path: 文件路径
        
        Returns:
            文档内容文本，失败返回None
        """
        try:
            file_ext = Path(file_path).suffix.lower()
            
            if file_ext == '.txt':
                return DocumentHandler._read_txt(file_path)
            elif file_ext == '.docx':
                return DocumentHandler._read_docx(file_path)
            elif file_ext == '.doc':
                return DocumentHandler._read_doc(file_path)
            elif file_ext == '.pdf':
                return DocumentHandler._read_pdf(file_path)
            elif file_ext in ['.md', '.markdown']:
                return DocumentHandler._read_markdown(file_path)
            elif file_ext == '.rtf':
                return DocumentHandler._read_rtf(file_path)
            elif file_ext == '.odt':
                return DocumentHandler._read_odt(file_path)
            elif file_ext in ['.html', '.htm']:
                return DocumentHandler._read_html(file_path)
            else:
                # 默认当作文本文件读取
                return DocumentHandler._read_txt(file_path)
        
        except Exception as e:
            print(f"[ERROR] 读取文档失败: {e}")
            return None
    
    @staticmethod
    def write_document(file_path: str, content: str) -> bool:
        """写入文档内容
        
        Args:
            file_path: 文件路径
            content: 要写入的内容
        
        Returns:
            成功返回True，失败返回False
        """
        try:
            file_ext = Path(file_path).suffix.lower()
            
            if file_ext == '.txt':
                return DocumentHandler._write_txt(file_path, content)
            elif file_ext == '.docx':
                return DocumentHandler._write_docx(file_path, content)
            elif file_ext == '.doc':
                # .doc格式写入转为.docx（推荐方式）
                print("[INFO] .doc格式写入将转换为.docx格式")
                docx_path = file_path.rsplit('.', 1)[0] + '.docx'
                return DocumentHandler._write_docx(docx_path, content)
            elif file_ext == '.pdf':
                return DocumentHandler._write_pdf(file_path, content)
            elif file_ext in ['.md', '.markdown']:
                return DocumentHandler._write_markdown(file_path, content)
            elif file_ext == '.rtf':
                return DocumentHandler._write_rtf(file_path, content)
            elif file_ext == '.odt':
                return DocumentHandler._write_odt(file_path, content)
            elif file_ext in ['.html', '.htm']:
                return DocumentHandler._write_html(file_path, content)
            else:
                # 默认当作文本文件写入
                return DocumentHandler._write_txt(file_path, content)
        
        except Exception as e:
            print(f"[ERROR] 写入文档失败: {e}")
            return False
    
    @staticmethod
    def _read_txt(file_path: str) -> str:
        """读取文本文件"""
        # 尝试多种编码
        encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        # 如果所有编码都失败，使用二进制模式读取并尝试解码
        with open(file_path, 'rb') as f:
            data = f.read()
            # 尝试检测编码
            try:
                return data.decode('utf-8')
            except:
                return data.decode('gbk', errors='ignore')
    
    @staticmethod
    def _write_txt(file_path: str, content: str) -> bool:
        """写入文本文件"""
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
            return True
        except Exception as e:
            print(f"[ERROR] 写入txt文件失败: {e}")
            return False
    
    @staticmethod
    def _read_docx(file_path: str) -> str:
        """读取Word文档（.docx）"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # 提取所有段落文本
            paragraphs = []
            for para in doc.paragraphs:
                paragraphs.append(para.text)
            
            # 用换行符连接
            return '\n'.join(paragraphs)
        
        except ImportError:
            raise ImportError("需要安装 python-docx: pip install python-docx")
        except Exception as e:
            print(f"[ERROR] 读取docx文件失败: {e}")
            raise
    
    @staticmethod
    def _write_docx(file_path: str, content: str) -> bool:
        """写入Word文档（.docx）"""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            
            doc = Document()
            
            # 设置默认字体和样式
            style = doc.styles['Normal']
            font = style.font
            font.name = '宋体'
            font.size = Pt(12)
            
            # 按段落分割内容
            paragraphs = content.split('\n')
            
            for para_text in paragraphs:
                if para_text.strip():  # 跳过空行
                    para = doc.add_paragraph(para_text)
                    # 设置段落格式
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    para_format = para.paragraph_format
                    para_format.line_spacing = 1.5  # 1.5倍行距
                else:
                    # 空行也添加，保持格式
                    doc.add_paragraph()
            
            # 保存文档
            doc.save(file_path)
            return True
        
        except ImportError:
            raise ImportError("需要安装 python-docx: pip install python-docx")
        except Exception as e:
            # 当 python-docx 的默认模板缺失时，尝试使用最小可用的 docx 结构写入
            err_msg = str(e)
            if "Package not found" in err_msg and "templates" in err_msg:
                try:
                    return DocumentHandler._write_docx_minimal_zip(file_path, content)
                except Exception as zip_e:
                    print(f"[ERROR] 通过最小ZIP写入docx仍失败: {zip_e}")
                    return False
            
            print(f"[ERROR] 写入docx文件失败: {e}")
            return False

    @staticmethod
    def _write_docx_minimal_zip(file_path: str, content: str) -> bool:
        """不依赖 python-docx 默认模板，使用最小Office Open XML结构写入docx。
        仅保留基本段落，适用于应急回退，尽量小改动保证功能可用。
        """
        # 基本的Content_Types定义
        content_types_xml = (
            """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>
"""
        )

        # 包级关系，指向word/document.xml
        rels_xml = (
            """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>
"""
        )

        # 将文本按段落输出为w:p/w:r/w:t结构
        paragraphs_xml_parts = []
        for line in content.split("\n"):
            if line.strip():
                paragraphs_xml_parts.append(
                    f"<w:p><w:r><w:t>{escape(line)}</w:t></w:r></w:p>"
                )
            else:
                paragraphs_xml_parts.append("<w:p/>")

        paragraphs_xml = "\n    ".join(paragraphs_xml_parts)

        # 最小文档结构
        document_xml = (
            f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    {paragraphs_xml}
    <w:sectPr/>
  </w:body>
</w:document>
"""
        )

        # 写入zip为docx
        # 确保输出目录存在
        out_dir = os.path.dirname(file_path)
        if out_dir:
            os.makedirs(out_dir, exist_ok=True)

        with zipfile.ZipFile(file_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("[Content_Types].xml", content_types_xml)
            zf.writestr("_rels/.rels", rels_xml)
            zf.writestr("word/document.xml", document_xml)

        return True
    
    @staticmethod
    def _read_doc(file_path: str) -> str:
        """读取Word文档（.doc）- 老格式
        
        尝试多种方法：
        1. 使用antiword（需要外部程序）
        2. 使用pywin32（仅Windows）
        3. 提示转换为docx
        """
        try:
            # 方法1: 尝试使用win32com（仅Windows）
            if sys.platform == 'win32':
                try:
                    import win32com.client
                    word = win32com.client.Dispatch("Word.Application")
                    word.Visible = False
                    doc = word.Documents.Open(os.path.abspath(file_path))
                    text = doc.Content.Text
                    doc.Close()
                    word.Quit()
                    return text
                except Exception as e:
                    print(f"[WARNING] 使用Word COM读取失败: {e}")
            
            # 方法2: 使用textract（跨平台，但需要额外依赖）
            try:
                import textract
                text = textract.process(file_path).decode('utf-8')
                return text
            except ImportError:
                print("[WARNING] textract未安装，无法读取.doc文件")
            except Exception as e:
                print(f"[WARNING] textract读取失败: {e}")
            
            # 如果都失败，提示用户转换
            raise ValueError(
                ".doc格式读取失败。建议：\n"
                "1. 使用Word打开并另存为.docx格式\n"
                "2. Windows系统可安装pywin32: pip install pywin32\n"
                "3. 或安装textract: pip install textract"
            )
        
        except Exception as e:
            print(f"[ERROR] 读取.doc文件失败: {e}")
            raise
    
    @staticmethod
    def _read_pdf(file_path: str) -> str:
        """读取PDF文档"""
        try:
            import PyPDF2
            
            text_content = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text_content.append(page.extract_text())
            
            return '\n'.join(text_content)
        
        except ImportError:
            raise ImportError("需要安装 PyPDF2: pip install PyPDF2")
        except Exception as e:
            print(f"[ERROR] 读取PDF文件失败: {e}")
            raise
    
    @staticmethod
    def _read_markdown(file_path: str) -> str:
        """读取Markdown文档"""
        # Markdown本质上就是文本文件
        return DocumentHandler._read_txt(file_path)
    
    @staticmethod
    def _read_rtf(file_path: str) -> str:
        """读取RTF文档"""
        try:
            from striprtf.striprtf import rtf_to_text
            
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                rtf_content = file.read()
            
            return rtf_to_text(rtf_content)
        
        except ImportError:
            raise ImportError("需要安装 striprtf: pip install striprtf")
        except Exception as e:
            print(f"[ERROR] 读取RTF文件失败: {e}")
            raise
    
    @staticmethod
    def _read_odt(file_path: str) -> str:
        """读取OpenDocument文本（.odt）"""
        try:
            from odf import text, teletype
            from odf.opendocument import load
            
            doc = load(file_path)
            all_text = []
            
            for paragraph in doc.getElementsByType(text.P):
                all_text.append(teletype.extractText(paragraph))
            
            return '\n'.join(all_text)
        
        except ImportError:
            raise ImportError("需要安装 odfpy: pip install odfpy")
        except Exception as e:
            print(f"[ERROR] 读取ODT文件失败: {e}")
            raise
    
    @staticmethod
    def _read_html(file_path: str) -> str:
        """读取HTML文档"""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8') as file:
                html_content = file.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # 移除script和style标签
            for script in soup(["script", "style"]):
                script.decompose()
            
            # 提取文本
            text = soup.get_text()
            
            # 清理多余空行
            lines = (line.strip() for line in text.splitlines())
            text = '\n'.join(line for line in lines if line)
            
            return text
        
        except ImportError:
            raise ImportError("需要安装 beautifulsoup4: pip install beautifulsoup4")
        except Exception as e:
            print(f"[ERROR] 读取HTML文件失败: {e}")
            raise
    
    @staticmethod
    def _write_pdf(file_path: str, content: str) -> bool:
        """写入PDF文档"""
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.lib.units import inch
            
            # 注册中文字体（尝试使用系统字体）
            try:
                # Windows
                if sys.platform == 'win32':
                    font_path = 'C:\\Windows\\Fonts\\simsun.ttc'
                    if os.path.exists(font_path):
                        pdfmetrics.registerFont(TTFont('SimSun', font_path))
                        font_name = 'SimSun'
                    else:
                        font_name = 'Helvetica'
                else:
                    font_name = 'Helvetica'
            except:
                font_name = 'Helvetica'
            
            # 创建PDF
            c = canvas.Canvas(file_path, pagesize=A4)
            width, height = A4
            
            # 设置字体
            c.setFont(font_name, 12)
            
            # 分页写入内容
            lines = content.split('\n')
            y_position = height - 1*inch
            line_height = 16
            
            for line in lines:
                if y_position < 1*inch:  # 需要换页
                    c.showPage()
                    c.setFont(font_name, 12)
                    y_position = height - 1*inch
                
                try:
                    c.drawString(1*inch, y_position, line)
                except:
                    # 如果字符编码有问题，跳过该行
                    pass
                
                y_position -= line_height
            
            c.save()
            return True
        
        except ImportError:
            raise ImportError("需要安装 reportlab: pip install reportlab")
        except Exception as e:
            print(f"[ERROR] 写入PDF文件失败: {e}")
            return False
    
    @staticmethod
    def _write_markdown(file_path: str, content: str) -> bool:
        """写入Markdown文档"""
        # Markdown本质上就是文本文件
        return DocumentHandler._write_txt(file_path, content)
    
    @staticmethod
    def _write_rtf(file_path: str, content: str) -> bool:
        """写入RTF文档"""
        try:
            from pyth.plugins.plaintext.reader import PlaintextReader
            from pyth.plugins.rtf15.writer import Rtf15Writer
            
            # 读取纯文本
            doc = PlaintextReader().read(content)
            
            # 写入RTF
            with open(file_path, 'wb') as file:
                Rtf15Writer.write(doc, file)
            
            return True
        
        except ImportError:
            # 如果pyth不可用，使用简单的RTF格式
            try:
                rtf_header = r"{\rtf1\ansi\deff0"
                rtf_footer = r"}"
                
                # 简单转义
                content_escaped = content.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')
                # 将换行转为RTF换行
                content_escaped = content_escaped.replace('\n', '\\par\n')
                
                rtf_content = f"{rtf_header}\n{content_escaped}\n{rtf_footer}"
                
                with open(file_path, 'w', encoding='utf-8') as file:
                    file.write(rtf_content)
                
                return True
            except Exception as e:
                print(f"[ERROR] 写入RTF文件失败: {e}")
                return False
        
        except Exception as e:
            print(f"[ERROR] 写入RTF文件失败: {e}")
            return False
    
    @staticmethod
    def _write_odt(file_path: str, content: str) -> bool:
        """写入OpenDocument文本（.odt）"""
        try:
            from odf.opendocument import OpenDocumentText
            from odf.text import P
            from odf.style import Style, TextProperties, ParagraphProperties
            from odf import style
            
            doc = OpenDocumentText()
            
            # 创建样式
            text_style = Style(name="TextBody", family="paragraph")
            text_style.addElement(ParagraphProperties(textalign="left"))
            text_style.addElement(TextProperties(fontsize="12pt", fontfamily="SimSun"))
            doc.styles.addElement(text_style)
            
            # 添加段落
            for line in content.split('\n'):
                p = P(stylename=text_style, text=line)
                doc.text.addElement(p)
            
            doc.save(file_path)
            return True
        
        except ImportError:
            raise ImportError("需要安装 odfpy: pip install odfpy")
        except Exception as e:
            print(f"[ERROR] 写入ODT文件失败: {e}")
            return False
    
    @staticmethod
    def _write_html(file_path: str, content: str) -> bool:
        """写入HTML文档"""
        try:
            # 创建简单的HTML文档
            html_template = """<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>文档</title>
    <style>
        body {{
            font-family: "Microsoft YaHei", "SimSun", sans-serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
        }}
        p {{
            margin-bottom: 1em;
        }}
    </style>
</head>
<body>
{content}
</body>
</html>"""
            
            # 将文本内容转换为HTML段落
            paragraphs = []
            for line in content.split('\n'):
                if line.strip():
                    # 转义HTML特殊字符
                    escaped_line = (line
                        .replace('&', '&amp;')
                        .replace('<', '&lt;')
                        .replace('>', '&gt;')
                        .replace('"', '&quot;'))
                    paragraphs.append(f"    <p>{escaped_line}</p>")
                else:
                    paragraphs.append("    <br>")
            
            html_content = html_template.format(content='\n'.join(paragraphs))
            
            with open(file_path, 'w', encoding='utf-8') as file:
                file.write(html_content)
            
            return True
        
        except Exception as e:
            print(f"[ERROR] 写入HTML文件失败: {e}")
            return False
    
    @staticmethod
    def create_new_document(file_path: str, template_content: str = "") -> bool:
        """创建新文档
        
        Args:
            file_path: 文件路径
            template_content: 模板内容（可选）
        
        Returns:
            成功返回True，失败返回False
        """
        try:
            # 确保目录存在
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            
            # 如果没有模板内容，使用默认模板
            if not template_content:
                template_content = "# 新建文档\n\n开始您的创作..."
            
            return DocumentHandler.write_document(file_path, template_content)
        
        except Exception as e:
            print(f"[ERROR] 创建新文档失败: {e}")
            return False
    
    @staticmethod
    def get_supported_formats() -> list:
        """获取支持的文件格式"""
        return [
            '.txt',     # 纯文本
            '.doc',     # Word旧格式
            '.docx',    # Word新格式
            '.pdf',     # PDF文档
            '.md',      # Markdown
            '.markdown',# Markdown
            '.rtf',     # RTF富文本
            '.odt',     # OpenDocument
            '.html',    # HTML
            '.htm',     # HTML
            '.epub'     # 电子书（ePub）
        ]
    
    @staticmethod
    def is_supported(file_path: str) -> bool:
        """检查文件格式是否支持"""
        ext = Path(file_path).suffix.lower()
        return ext in DocumentHandler.get_supported_formats()
    
    @staticmethod
    def get_format_description(ext: str) -> str:
        """获取文件格式的描述"""
        descriptions = {
            '.txt': '纯文本文件',
            '.doc': 'Word文档（旧格式）',
            '.docx': 'Word文档',
            '.pdf': 'PDF文档',
            '.md': 'Markdown文档',
            '.markdown': 'Markdown文档',
            '.rtf': 'RTF富文本',
            '.odt': 'OpenDocument文本',
            '.html': 'HTML网页',
            '.htm': 'HTML网页',
            '.epub': 'ePub电子书'
        }
        return descriptions.get(ext.lower(), '未知格式')

